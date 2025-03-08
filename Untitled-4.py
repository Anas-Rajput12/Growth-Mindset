import streamlit as st
import pandas as pd
import json
import io
import PyPDF2
from docx import Document

# Set Page Title
st.set_page_config(page_title="Universal Data Sweeper", page_icon="🧹", layout="centered")

# Title
st.title("🧹 Universal Data Sweeper - Clean Any File!")
st.write("Upload a **CSV, Excel, TXT, JSON, PDF, or Word** file, and we'll clean it for you!")

# File Upload
uploaded_file = st.file_uploader("📂 Upload Your File", type=["csv", "xlsx", "txt", "json", "pdf", "docx"])

# Data Cleaning Function
def clean_text(text):
    """Removes special characters & trims spaces."""
    return ''.join(e for e in text if e.isalnum() or e.isspace()).strip()

def clean_data(df):
    """Cleans the dataframe by removing duplicates, empty rows, and special characters."""
    df.drop_duplicates(inplace=True)  # Remove Duplicates
    df.dropna(inplace=True)  # Remove Empty Rows
    df = df.applymap(lambda x: clean_text(str(x)) if isinstance(x, str) else x)  # Clean text
    return df

if uploaded_file:
    try:
        file_type = uploaded_file.name.split(".")[-1]
        df = None  # Initialize dataframe

        # ✅ Load & Process Data Based on File Type
        if file_type == "csv":
            df = pd.read_csv(uploaded_file)
        elif file_type == "xlsx":
            df = pd.read_excel(uploaded_file)
        elif file_type == "txt":
            text_data = uploaded_file.getvalue().decode("utf-8").splitlines()
            df = pd.DataFrame(text_data, columns=["Text"])
        elif file_type == "json":
            json_data = json.load(uploaded_file)
            df = pd.DataFrame(json_data)
        elif file_type == "pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_data = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            df = pd.DataFrame(text_data.split("\n"), columns=["Text"])
        elif file_type == "docx":
            doc = Document(uploaded_file)
            text_data = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            df = pd.DataFrame(text_data.split("\n"), columns=["Text"])

        # Display Original Data
        st.subheader("📊 Original Data Preview")
        st.dataframe(df.head())

        # Clean Data
        cleaned_df = clean_data(df)
        st.subheader("✅ Cleaned Data Preview")
        st.dataframe(cleaned_df.head())

        # ✅ Create Downloadable Cleaned File
        output = io.BytesIO()
        file_name = f"cleaned_data.{file_type}"

        if file_type == "csv":
            cleaned_df.to_csv(output, index=False, encoding="utf-8-sig")  # Fix encoding issue
            mime_type = "text/csv"
        elif file_type == "xlsx":
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                cleaned_df.to_excel(writer, index=False)
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif file_type == "json":
            json.dump(cleaned_df.to_dict(orient="records"), output, indent=4)
            mime_type = "application/json"
        elif file_type == "txt":
            output.write("\n".join(cleaned_df["Text"]).encode("utf-8"))
            mime_type = "text/plain"
        elif file_type == "pdf":  # ❌ Instead of PDF, save as TXT to avoid errors
            output.write("\n".join(cleaned_df["Text"]).encode("utf-8"))
            mime_type = "text/plain"
            file_name = "cleaned_data.txt"
        elif file_type == "docx":
            new_doc = Document()
            for line in cleaned_df["Text"]:
                new_doc.add_paragraph(line)
            new_doc.save(output)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        output.seek(0)

        # ✅ Ensure File is Downloadable
        st.download_button(
            label="⬇️ Download Cleaned File",
            data=output.getvalue(),  # Ensure correct file content
            file_name=file_name,
            mime=mime_type
        )

    except Exception as e:
        st.error(f"❌ Error: {str(e)}")

# Footer
st.write("✨ Built with ❤️ using Streamlit")
